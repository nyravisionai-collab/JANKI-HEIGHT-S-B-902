# -*- coding: utf-8 -*-
"""Generate the Janki Heights B-901 Flat-901 electrical quotation DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1F, 0x3B, 0x73)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# ---------- base styles ----------
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

for sec in doc.sections:
    sec.top_margin = Cm(1.4); sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)

def para(text='', size=10, bold=False, italic=False, color=None, align=None, space_after=4, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def heading(text, size=13, color=BLUE, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, size=9, align=None, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color

# =====================================================================
# COVER
# =====================================================================
para('GHANSHYAM ELECTRICALS', 20, True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('441-The Galleria, Yogi Chowk, Surat, Gujarat - 395006', 10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para('Mobile: +91 92655-60638   |   Email: ghanshyamelectricals101@gmail.com', 10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

para('ELECTRICAL WORKS QUOTATION', 17, True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para('Prepared from the approved electrical layout drawings (Ceiling Plan, Light Dimension Layout & Electric Wiring Layout)', 10.5, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# project info table
info = doc.add_table(rows=4, cols=4)
info.style = 'Table Grid'
info.alignment = WD_TABLE_ALIGNMENT.CENTER
rows_data = [
    ('Project', 'JANKI HEIGHT\u2019S, B-Building, 9th Floor, Flat No. 901', 'Client Name', '________________________'),
    ('Scope', 'Interior Electrical Works \u2013 Flat 901 (as per drawing)', 'Contact No.', '________________________'),
    ('Drawing Ref.', 'BULD-B 901_JANKI HEIGHTS_electric plan.dwg', 'Site Address', '________________________'),
    ('Quotation No.', 'GE-874 (Rev. 1)   |   Date: 15/08/2026   |   Valid: 10 days', 'Ceiling Area', '458.64 sq.ft + 46.79 sq.ft (toilet)'),
]
for i, (a, b, c, d) in enumerate(rows_data):
    cells = info.rows[i].cells
    set_cell(cells[0], a, bold=True, size=9); shade(cells[0], 'E8EDF5')
    set_cell(cells[1], b, size=9)
    set_cell(cells[2], c, bold=True, size=9); shade(cells[2], 'E8EDF5')
    set_cell(cells[3], d, size=9)
for r in info.rows:
    r.cells[0].width = Cm(2.8); r.cells[2].width = Cm(2.8)
    r.cells[1].width = Cm(6.4); r.cells[3].width = Cm(4.5)

para('', 4, space_after=2)
para('Note: This quotation is prepared after a systematic reading of the client-supplied electrical layout. The drawing contains three sheets of Flat 901 \u2013 (1) Ceiling Plan (light/fan/exhaust/geyser points), (2) Light Dimension Layout (same ceiling scope, dimensioned), and (3) Electric Wiring Layout (switch boards, plug points, A.C. points and circuit annotations). Quantities below are counted point-by-point from these sheets.', 9, italic=True, color=GRAY, space_after=8)

# =====================================================================
# SECTION 1 - LAYOUT UNDERSTANDING (room-wise schedule)
# =====================================================================
heading('1. ELECTRICAL LAYOUT \u2013 ROOM-WISE POINT SCHEDULE (FLAT 901)')
para('Summary of every electrical point counted from the drawing, room by room:', 9.5, space_after=4)

sched = doc.add_table(rows=12, cols=11)
sched.style = 'Table Grid'
hdr = ['Room', 'Ceiling\nLight', 'Fan', 'Wall\nLight', 'Stripe\nLight', '5A\nSocket', '15A\nPoint', 'Switch\nBoard', 'Geyser', 'Exhaust', 'Others']
rows = [
    ('Living', '4', '2', '1', '1', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', 'Pooja light \u2013 1, A.C. O.D.U. \u2013 1'),
    ('Kitchen', '6', '1', '\u2013', '\u2013', '\u2013', '\u2013', '1', '\u2013', '\u2013', 'Chimney \u2013 1, Water filter \u2013 1'),
    ('Bedroom 1', '4', '1', '\u2013', '1', '3', '1', '2*', '\u2013', '\u2013', 'Split A.C. point \u2013 1, O.D.U. \u2013 1'),
    ('Bedroom 2', '4', '1', '1', '2', '2', '1', '2', '\u2013', '\u2013', '\u2013'),
    ('A. Toilet', '1', '\u2013', '1', '\u2013', '\u2013', '\u2013', '3', '1', '1', '\u2013'),
    ('C. Toilet', '2', '\u2013', '1', '\u2013', '\u2013', '2', '3', '1', '1', '\u2013'),
    ('Passage', '1', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', 'Bell \u2013 1'),
    ('Entry', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '2', '\u2013', '\u2013', '\u2013'),
    ('Wash / Utility', '\u2013', '\u2013', '1', '\u2013', '\u2013', '2', '2', '\u2013', '\u2013', 'Internet point \u2013 1'),
    ('Store / Utility wall', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '\u2013', '1', '\u2013', '\u2013'),
    ('TOTAL', '22', '5', '4', '4', '5', '6', '15', '3', '2', '9 misc. points'),
]
for j, h in enumerate(hdr):
    set_cell(sched.rows[0].cells[j], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(sched.rows[0].cells[j], '1F3B73')
    sched.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for i, r in enumerate(rows, start=1):
    for j, v in enumerate(r):
        set_cell(sched.rows[i].cells[j], v, bold=(r[0] == 'TOTAL'), size=8,
                 align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else None)
    if r[0] == 'TOTAL':
        for c in sched.rows[i].cells: shade(c, 'F2E3C8')
para('* Bedroom switch boards include bed-panel boards at 3\u2032-0\u2033 / 2\u2032-6\u2033 height as per the drawing\u2019s switch-board height schedule.', 8, italic=True, color=GRAY, space_after=4)
para('Additional scope noted on the drawing: 2 Nos. A.C. Outdoor Units (living + bedroom, above/below), pooja-light point in living furniture, chimney & water-filter points in kitchen, and one-way / two-way ceiling circuit lines (~99 m drawn length). Ceiling areas: living etc. 458.64 sq.ft + toilet 46.79 sq.ft (total 505.43 sq.ft \u2248 47.0 m\u00b2).', 9, space_after=8)

# =====================================================================
# SECTION 2 - BOQ
# =====================================================================
heading('2. PRICING \u2013 BOQ (QUANTITIES FROM DRAWING, RATES INDICATIVE)')
para('Rates include standard/ISI-marked material and labour. GST extra. Final quantities shall be confirmed on-site before commencement.', 9, italic=True, color=GRAY, space_after=4)

boq = doc.add_table(rows=1, cols=7)
boq.style = 'Table Grid'
hdrs = ['Sr.', 'Description of Work', 'Unit', 'Qty.', 'Rate \u2013 Material + Labour (\u20b9)', 'Amount (\u20b9)', 'Labour Only (\u20b9)']
for j, h in enumerate(hdrs):
    set_cell(boq.rows[0].cells[j], h, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(boq.rows[0].cells[j], '1F3B73')
    boq.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

items = [
    # (section, sr, desc, unit, qty, rate_ml, rate_lo)
    ('A. Distribution & Panel Work', 'A1', 'Distribution Board (SDB) with MCB + RCCB set, supply & installation', 'Set', 1, 12500, 2000),
    ('A. Distribution & Panel Work', 'A2', 'MCB / RCCB fitting & circuit connection', 'Per pole', 12, 850, 150),
    ('A. Distribution & Panel Work', 'A3', 'Switch board installation & circuit wiring (standard, 4\u2032-6\u2033 height)', 'Per board', 11, 1250, 450),
    ('A. Distribution & Panel Work', 'A4', 'Switch board @ 3\u2032-0\u2033 / 2\u2032-6\u2033 (bed panel / geyser boards)', 'Per board', 4, 1250, 450),
    ('A. Distribution & Panel Work', 'A5', 'Panel labelling & circuit identification', 'Lot', 1, 1500, 500),
    ('B. Light & Power Wiring', 'B1', 'Ceiling light point (false-ceiling / downlight) \u2013 concealed wiring', 'Per point', 22, 950, 300),
    ('B. Light & Power Wiring', 'B2', 'Wall light point at 7\u2032-0\u2033', 'Per point', 4, 1100, 350),
    ('B. Light & Power Wiring', 'B3', 'Stripe / profile LED light point', 'Per point', 4, 1300, 400),
    ('B. Light & Power Wiring', 'B4', 'Pooja light point (in furniture)', 'Per point', 1, 950, 300),
    ('B. Light & Power Wiring', 'B5', 'Fan point with regulator connection', 'Per point', 5, 1450, 400),
    ('B. Light & Power Wiring', 'B6', '5A socket (plug) point', 'Per point', 5, 900, 300),
    ('B. Light & Power Wiring', 'B7', '15A power point', 'Per point', 6, 1200, 350),
    ('B. Light & Power Wiring', 'B8', 'Split A.C. point (indoor unit, 4 sq.mm wiring + isolator)', 'Per point', 2, 2600, 600),
    ('B. Light & Power Wiring', 'B9', 'A.C. outdoor unit (O.D.U.) point', 'Per point', 2, 1800, 450),
    ('B. Light & Power Wiring', 'B10', 'Geyser / water heater point', 'Per point', 3, 2100, 450),
    ('B. Light & Power Wiring', 'B11', 'Exhaust fan point', 'Per point', 2, 1200, 350),
    ('B. Light & Power Wiring', 'B12', 'Call bell point', 'Per point', 1, 1100, 350),
    ('B. Light & Power Wiring', 'B13', 'Chimney point', 'Per point', 1, 1800, 450),
    ('B. Light & Power Wiring', 'B14', 'Water filter / RO point', 'Per point', 1, 1100, 350),
    ('C. Cabling & Conduit', 'C1', 'Main power cable laying (floor riser to flat DB)', 'Per meter', 15, 420, 80),
    ('C. Cabling & Conduit', 'C2', 'Sub-circuit cable laying \u2013 2.5 / 4 sq.mm FR (approx.)', 'Per meter', 750, 72, 20),
    ('C. Cabling & Conduit', 'C3', 'ISI concealed conduit fitting & laying (approx.)', 'Per meter', 700, 48, 18),
    ('C. Cabling & Conduit', 'C4', 'Junction box / pull box fitting', 'Per unit', 8, 250, 60),
    ('C. Cabling & Conduit', 'C5', 'Earthing cable laying & connection', 'Per meter', 10, 180, 30),
    ('D. Data & Communication', 'D1', 'CAT 6 network cable laying (approx.)', 'Per meter', 40, 38, 12),
    ('D. Data & Communication', 'D2', 'Network patch panel / jack installation', 'Per unit', 1, 1200, 300),
    ('G. Special & Miscellaneous', 'G1', 'Testing, commissioning & handover', 'Lump sum', 1, 4500, 1500),
    ('G. Special & Miscellaneous', 'G2', 'Any other work as per site requirement', 'As agreed', 0, 0, 0),
]

sec_total = {}
grand_ml = grand_lo = 0
cur_sec = None
for sec, sr, desc, unit, qty, r_ml, r_lo in items:
    if sec != cur_sec:
        cur_sec = sec
        row = boq.add_row()
        set_cell(row.cells[0], sec, bold=True, size=9)
        for c in row.cells[1:]: set_cell(c, '', size=9)
        for c in row.cells: shade(c, 'DCE4F0')
        sec_total[sec] = [0, 0]
    row = boq.add_row()
    set_cell(row.cells[0], sr, size=9)
    set_cell(row.cells[1], desc, size=9)
    set_cell(row.cells[2], unit, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[3], f'{qty:,}' if qty else '\u2013', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[4], f'{r_ml:,.0f}' if qty else '\u2013', size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    amt = qty * r_ml
    set_cell(row.cells[5], f'{amt:,.0f}' if qty else '\u2013', size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(row.cells[6], f'{r_lo:,.0f}' if qty else '\u2013', size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    sec_total[sec][0] += amt
    sec_total[sec][1] += qty * r_lo
    grand_ml += amt; grand_lo += qty * r_lo

# totals row
row = boq.add_row()
set_cell(row.cells[0], '', size=9); set_cell(row.cells[1], 'GRAND TOTAL', bold=True, size=10)
set_cell(row.cells[2], '', size=9); set_cell(row.cells[3], '', size=9); set_cell(row.cells[4], '', size=9)
set_cell(row.cells[5], f'\u20b9 {grand_ml:,.0f}', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell(row.cells[6], f'\u20b9 {grand_lo:,.0f}', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
for c in row.cells: shade(c, 'F2E3C8')

widths = [Cm(3.4), Cm(8.6), Cm(1.7), Cm(1.3), Cm(2.2), Cm(1.9), Cm(2.0)]
for r in boq.rows:
    for j, c in enumerate(r.cells):
        try: c.width = widths[j]
        except Exception: pass

para('', 4, space_after=2)
para(f'Grand Total \u2013 With Material & Labour:  \u20b9 {grand_ml:,.0f}   |   Grand Total \u2013 Labour Only:  \u20b9 {grand_lo:,.0f}', 11, True, space_after=2)
para('GST @ 18% and any other applicable taxes extra. Rates are indicative market rates for standard/ISI material; contractor\u2019s final rates to apply before issue.', 9, italic=True, color=GRAY, space_after=8)

# =====================================================================
# SECTION 3 - ASSUMPTIONS
# =====================================================================
heading('3. ASSUMPTIONS & NOTES ON QUANTITIES')
for line in [
    '1. Quantities are counted point-by-point from the client-supplied drawing \u201cBULD-B 901_JANKI HEIGHTS_electric plan.dwg\u201d (Ceiling Plan, Light Dimension Layout and Electric Wiring Layout of Flat 901, B-Building, 9th Floor, Janki Height\u2019s).',
    '2. The ceiling plan and light-dimension layout represent the same physical points; they have been counted once (22 ceiling lights, 5 fans, 3 geyser points).',
    '3. Conduit & sub-circuit cable quantities are estimated on a per-point basis (~11\u201312 m per point) plus ~99 m of one-way/two-way ceiling circuit lines measured on the drawing. Final running-meter quantities to be measured on site.',
    '4. Fire alarm, CCTV, access control, audio-visual and automation items are not shown on this interior electrical drawing and are therefore quoted as \u201cNot Applicable \u2013 on request\u201d.',
    '5. Earth pit is assumed existing; if a new pit is required it shall be quoted separately.',
    '6. The drawing is an interior-design-purpose plan (Architects: Ar. Tirth Radadiya / Ar. Raj Modi). Any deviation from the approved layout requires written approval.',
]:
    para(line, 9, space_after=2)

# =====================================================================
# SECTION 4 - TERMS
# =====================================================================
heading('4. TERMS & CONDITIONS (CONDENSED \u2013 FULL VERSION AS PER GE-874)')
for line in [
    'Payment: 50% advance before commencement; 25% on 50% work completion; balance 25% within 3 days of completion. 2% per month interest on overdue amounts.',
    'Materials: ISI-marked/standard quality only. Material substitution only with prior written client approval.',
    'Warranty: 1 year on workmanship; material warranty as per manufacturer.',
    'Work hours: Mon\u2013Sat, 9:00 AM\u20136:00 PM. As-built drawings, test reports and completion certificate on handover.',
    'Confidentiality: All drawings, data and pricing remain confidential under the NDA (valid during project + 2 years thereafter).',
    'Jurisdiction: Indian Electricity Act 2003 / IS / IEC standards; disputes subject to Surat, Gujarat courts.',
]:
    para(line, 9, space_after=2)

para('', 6, space_after=2)
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
set_cell(t.rows[0].cells[0], 'Client / Customer\nSignature: ______________\nName & Date: ______________', size=9)
set_cell(t.rows[0].cells[1], 'Quotation Accepted\n(Yes / No): ______________\nDate: ______________', size=9)
set_cell(t.rows[0].cells[2], 'For Ghanshyam Electricals\nSignature: ______________\nName & Date: ______________', size=9)
for c in t.rows[0].cells:
    c.width = Cm(6.3)

para('', 6, space_after=2)
heading('ANNEXURE A \u2013 ELECTRICAL LAYOUT REFERENCE (EXTRACTED FROM CLIENT DWG)', 11)
para('The figure below reproduces the electrical layout of Flat 901 (three sheets + legend) as read from the supplied DWG file.', 9, italic=True, color=GRAY, space_after=4)
doc.add_picture('/tmp/dwgconv/electrical_plan_overview.png', width=Cm(18.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para('This document is confidential and intended solely for the named client. Unauthorized reproduction or distribution is strictly prohibited.', 8, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)

out = '/home/user/JANKI-HEIGHT-S-B-902/GE_Quotation_JankiHeights_B901_Flat901.docx'
doc.save(out)
print('saved', out)
print('grand material+labour:', grand_ml, ' labour only:', grand_lo)
